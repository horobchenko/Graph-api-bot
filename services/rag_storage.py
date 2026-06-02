"""Document ingestion and semantic retrieval using ChromaDB and OpenAI embeddings."""

from pathlib import Path
from typing import Any

from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument

from config.settings import CHROMA_DIR, OPENAI_API_KEY


class RAGStorage:
    """Simple RAG indexer and retriever for local documents."""

    def __init__(self, persist_directory: str | Path = CHROMA_DIR, collection_name: str = "real_estate_docs") -> None:
        self.persist_directory = Path(persist_directory)
        self.persist_directory.mkdir(parents=True, exist_ok=True)
        self.collection_name = collection_name
        self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small", api_key=OPENAI_API_KEY)
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=120)

    @staticmethod
    def _read_text(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".txt":
            return path.read_text(encoding="utf-8")
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".docx":
            document = DocxDocument(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        if suffix in {".md", ".rst"}:
            return path.read_text(encoding="utf-8")
        raise ValueError(f"Unsupported file type: {suffix}")

    def index_files(self, file_paths: list[str | Path]) -> None:
        """Load files, split them into chunks, and store them in ChromaDB."""
        documents: list[Document] = []

        for file_path in file_paths:
            path = Path(file_path)
            if not path.exists():
                continue
            text = self._read_text(path)
            if not text.strip():
                continue

            chunks = self.text_splitter.split_text(text)
            for index, chunk in enumerate(chunks):
                documents.append(
                    Document(
                        page_content=chunk,
                        metadata={
                            "source": str(path),
                            "chunk_index": index,
                            "file_name": path.name,
                        },
                    )
                )

        if not documents:
            return

        vector_store = Chroma.from_documents(
            documents=documents,
            embedding=self.embeddings,
            persist_directory=str(self.persist_directory),
            collection_name=self.collection_name,
        )
        vector_store.persist()

    def semantic_search(self, query: str, k: int = 4) -> list[dict[str, Any]]:
        """Return relevant chunks for a semantic query."""
        vector_store = Chroma(
            persist_directory=str(self.persist_directory),
            embedding_function=self.embeddings,
            collection_name=self.collection_name,
        )
        results = vector_store.similarity_search_with_score(query=query, k=k)
        return [
            {
                "content": item[0].page_content,
                "metadata": item[0].metadata,
                "score": item[1],
            }
            for item in results
        ]
